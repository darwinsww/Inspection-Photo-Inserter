"""Word document creation and orientation-aware photo layout.

The output document is generated programmatically from the Markdown template
definition (spec 3.6): a cover page and an "Areas" index (both photo-free),
followed by the fixed list of titles, each on its own page. Photos are inserted
only under the title matching their source folder name.

Layout rules (see specification sections 7 & 8):
  * Portrait and landscape images never share a row.
  * Two images per row, minimizing single-image rows by pairing same-orientation
    photos within a folder (spec 7.1).
  * Row pairing never crosses folder boundaries (spec 7.2).
  * Same-named folders from different paths share a title but are separated by a
    thick red divider line (spec 3.5).
  * Image width is derived from the page size and margins, never hardcoded.
"""
from __future__ import annotations

import logging
from collections import OrderedDict
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Dict, List, Optional, Sequence, Tuple

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Length, Pt, RGBColor
from lxml import etree

from image_processor import Orientation, ProcessedImage
from template_loader import Template, normalize_title

logger = logging.getLogger(__name__)

# Callback signature: (images_done, images_total) -> None
ProgressCallback = Callable[[int, int], None]

# Plain black for headings/text (overrides the blue heading-style default).
BLACK = RGBColor(0x00, 0x00, 0x00)


@dataclass
class LayoutConfig:
    """Configuration for how photos are arranged on the page."""

    images_per_row: int = 2
    column_gap_cm: float = 0.5
    row_space_after_pt: int = 3


@dataclass
class FolderPhotos:
    """Processed images that all came from one physical folder.

    ``section_name`` is the immediate parent folder name used to group
    same-named folders into a single document section.
    """

    section_name: str
    folder: Path
    images: List[ProcessedImage]


def compact_rows(
    images: Sequence[ProcessedImage], images_per_row: int
) -> List[List[ProcessedImage]]:
    """Pack a single folder's images into rows, minimizing single-image rows.

    Same-orientation images are paired together even when they are not adjacent
    (spec 7.1): a lone landscape/portrait waits for the next image of the same
    orientation and they are emitted together. Any images still unpaired after
    the whole folder is scanned are emitted as final single/partial rows.

    This operates on one folder only; callers must not mix folders here so that
    pairing never crosses a folder boundary (spec 7.2).
    """
    rows: List[List[ProcessedImage]] = []
    buckets: Dict[Orientation, List[ProcessedImage]] = {
        Orientation.LANDSCAPE: [],
        Orientation.PORTRAIT: [],
    }

    for image in images:
        bucket = buckets[image.orientation]
        bucket.append(image)
        if len(bucket) == images_per_row:
            rows.append(list(bucket))
            bucket.clear()

    # Emit leftovers deterministically: landscape first, then portrait.
    for orientation in (Orientation.LANDSCAPE, Orientation.PORTRAIT):
        if buckets[orientation]:
            rows.append(list(buckets[orientation]))
    return rows


class WordProcessor:
    """Builds the output Word document from the template and processed images."""

    def __init__(self, layout: LayoutConfig | None = None) -> None:
        self.layout = layout or LayoutConfig()

    def build_report(
        self,
        template: Template,
        folders: Sequence[FolderPhotos],
        progress: ProgressCallback | None = None,
    ) -> Tuple[DocumentType, List[str]]:
        """Build the output document from *template* and insert *folders*.

        The document is generated programmatically from the template definition
        (spec 3.6): a cover page and an "Areas" index (both photo-free), followed
        by the fixed list of titles, each starting on its own page. Photos are
        placed only under the title whose name matches their source folder name
        (case-insensitive); the titles keep their template order and are never
        changed, and titles with no matching photos are still emitted empty.

        Returns the document and the list of folder names that matched no title
        (skipped, for the summary). Temporary image files are deleted as soon as
        they are embedded.
        """
        document = Document()
        self._force_theme_body_font(document)
        image_width = self._image_width(document)

        # Map each title to the folders that match it, preserving folder order.
        title_keys = {normalize_title(title) for title in template.titles}
        matched: "OrderedDict[str, List[FolderPhotos]]" = OrderedDict()
        unmatched: List[str] = []
        for folder in folders:
            key = normalize_title(folder.section_name)
            if key in title_keys:
                matched.setdefault(key, []).append(folder)
            else:
                unmatched.append(folder.section_name)

        total = sum(
            len(folder.images) for group in matched.values() for folder in group
        )
        done = 0

        self._add_cover_page(document, template)
        self._add_page_break(document)
        self._add_areas_index(document, template)

        # Emit every title in template order, even when it has no photos.
        for number, title in enumerate(template.titles, start=1):
            self._add_page_break(document)
            self._add_title(document, number, title)
            for index, folder in enumerate(matched.get(normalize_title(title), [])):
                # Same title, different physical folder: separate with a rule.
                if index > 0:
                    self._add_separator(document)
                for row in compact_rows(folder.images, self.layout.images_per_row):
                    self._add_photo_row(document, row, image_width)
                    done += len(row)
                    for image in row:
                        # Safe: python-docx has already embedded the picture.
                        try:
                            image.temp_path.unlink(missing_ok=True)
                        except OSError:
                            pass
                    if progress:
                        progress(done, total)

        if unmatched:
            logger.info("Folders with no matching title: %s", ", ".join(unmatched))
        return document, unmatched

    def save(self, document: DocumentType, output_path: Path) -> None:
        """Save the document to *output_path* (never the original template).

        A ``.docx`` extension is enforced so the file is recognized as a Word
        document (correct icon and default application) by the OS.
        """
        if output_path.suffix.lower() != ".docx":
            output_path = output_path.with_suffix(".docx")
        document.save(str(output_path))
        logger.info("Saved document: %s", output_path)

    # Internal helpers ----------------------------------------------------

    def _image_width(self, document: DocumentType) -> Length:
        """Compute per-image width from page size and margins.

        available_width = page_width - left_margin - right_margin
        image_width = (available_width - total_gap) / images_per_row
        """
        section = document.sections[0]
        available = section.page_width - section.left_margin - section.right_margin
        per_row = self.layout.images_per_row
        total_gap = Cm(self.layout.column_gap_cm) * (per_row - 1)
        width = int((available - total_gap) / per_row)
        return Emu(width)

    def _add_photo_row(
        self,
        document: DocumentType,
        row: Sequence[ProcessedImage],
        image_width: Length,
    ) -> None:
        """Add one row of images using a borderless table.

        Full rows are centered; a row with a single photo is left-aligned so the
        lone image sits on the left instead of the middle of the line (spec 8).
        """
        gap = Cm(self.layout.column_gap_cm)
        count = len(row)
        # Interleave narrow gap columns between image columns.
        column_count = count * 2 - 1
        table = document.add_table(rows=1, cols=column_count)
        single = count == 1
        table.alignment = (
            WD_TABLE_ALIGNMENT.LEFT if single else WD_TABLE_ALIGNMENT.CENTER
        )
        table.autofit = False

        cells = table.rows[0].cells
        for index, image in enumerate(row):
            image_cell = cells[index * 2]
            image_cell.width = image_width
            paragraph = image_cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if single else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = paragraph.add_run()
            # Width fixes the size; height follows to preserve aspect ratio.
            run.add_picture(str(image.temp_path), width=image_width)

            if index < count - 1:
                cells[index * 2 + 1].width = gap

        self._add_spacer(document)

    def _add_spacer(self, document: DocumentType) -> None:
        """Add a minimal paragraph that both spaces rows and separates tables.

        The paragraph prevents adjacent tables from merging and provides the
        small vertical gap after each row.
        """
        paragraph = document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(self.layout.row_space_after_pt)
        run = paragraph.add_run()
        run.font.size = Pt(1)

    def _add_cover_page(self, document: DocumentType, template: Template) -> None:
        """Generate the cover page (content only, no photos) (spec 3.6)."""
        if template.cover_title:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._body_font(paragraph.add_run(template.cover_title), 16, bold=True)
        # An empty line below the report title.
        document.add_paragraph()
        for line in template.cover_lines:
            self._add_cover_field(document, line)

    def _add_cover_field(self, document: DocumentType, text: str) -> None:
        """Render one cover field: a bold ``Label:`` plus a non-bold space.

        A blank template line renders as an empty paragraph. The label (up to and
        including the colon) is bold; the trailing space is not, so anything the
        user types after it stays un-bold. Uses Calibri (Body) 12pt.
        """
        stripped = text.strip()
        if not stripped:
            document.add_paragraph()
            return
        paragraph = document.add_paragraph()
        if ":" in stripped:
            label = stripped.split(":", 1)[0]
            self._body_font(paragraph.add_run(f"{label}:"), 12, bold=True)
            self._body_font(paragraph.add_run(" "), 12, bold=False)
        else:
            self._body_font(paragraph.add_run(stripped), 12, bold=True)

    def _add_areas_index(self, document: DocumentType, template: Template) -> None:
        """Generate the "Areas" index page (content only, no photos) (spec 3.6)."""
        heading = self._add_heading(document, template.areas_heading or "Areas", level=1)
        self._apply_font(heading, 22)  # Calibri, 22pt, black
        heading.paragraph_format.line_spacing = 2.41  # multiple
        for area in template.areas:
            # No numbering; Calibri 12pt; 1.5 line spacing (spec 3.6).
            paragraph = document.add_paragraph()
            self._body_font(paragraph.add_run(area), 12)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def _add_title(self, document: DocumentType, number: int, name: str) -> None:
        """Add a numbered photo-section title as Heading 3 (spec 3.6)."""
        heading = self._add_heading(document, f"{number}) {name}", level=3)
        self._apply_font(heading, 16)  # Calibri, 16pt, black
        heading.paragraph_format.line_spacing = 1.73  # multiple

    def _add_heading(self, document: DocumentType, text: str, level: int):
        """Add a heading paragraph, falling back to plain text if style missing."""
        try:
            return document.add_heading(text, level=level)
        except KeyError:
            paragraph = document.add_paragraph()
            paragraph.add_run(text)
            return paragraph

    def _apply_font(self, paragraph, size_pt: int) -> None:
        """Apply Calibri (Body) at *size_pt*, black, to every run in *paragraph*."""
        for run in paragraph.runs:
            self._body_font(run, size_pt)

    def _body_font(self, run, size_pt: int, bold: Optional[bool] = None) -> None:
        """Set a run to the Calibri (Body) theme font, *size_pt*, black.

        Runs reference the theme body font (``minorHAnsi``) so Word shows
        "Calibri (Body)"; :meth:`_force_theme_body_font` forces that theme font
        to Calibri (the default theme may otherwise resolve it to Cambria).
        """
        run.font.size = Pt(size_pt)
        run.font.color.rgb = BLACK
        if bold is not None:
            run.bold = bold
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        # Reference the theme body font, and drop any explicit face so it wins.
        rfonts.set(qn("w:asciiTheme"), "minorHAnsi")
        rfonts.set(qn("w:hAnsiTheme"), "minorHAnsi")
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            if rfonts.get(qn(attr)) is not None:
                del rfonts.attrib[qn(attr)]

    def _force_theme_body_font(
        self, document: DocumentType, name: str = "Calibri"
    ) -> None:
        """Force the theme's major/minor Latin fonts to *name* (Calibri).

        This makes "(Body)"/"(Headings)" theme references resolve to Calibri, so
        text set via :meth:`_body_font` renders as Calibri and displays as
        "Calibri (Body)" in Word, rather than the theme's default (e.g. Cambria).
        """
        a = "http://schemas.openxmlformats.org/drawingml/2006/main"
        try:
            for rel in document.part.rels.values():
                if "theme" not in rel.reltype:
                    continue
                part = rel.target_part
                root = etree.fromstring(part.blob)
                for scheme in ("majorFont", "minorFont"):
                    latin = root.find(
                        f"{{{a}}}themeElements/{{{a}}}fontScheme/"
                        f"{{{a}}}{scheme}/{{{a}}}latin"
                    )
                    if latin is not None:
                        latin.set("typeface", name)
                part._blob = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )
                break
        except Exception:  # pragma: no cover - best-effort theme tweak
            logger.debug("Could not adjust theme fonts", exc_info=True)

    def _add_page_break(self, document: DocumentType) -> None:
        """Start a new page so each title begins on its own page (spec 3.6)."""
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def _add_separator(self, document: DocumentType) -> None:
        """Add a thick red divider between same-named folders (spec 3.5)."""
        paragraph = document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(6)
        p_pr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "36")  # ~4.5pt: a clearly thick rule
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "FF0000")  # red
        borders.append(bottom)
        p_pr.append(borders)

